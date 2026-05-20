#!/usr/bin/env python3
"""
研究报告生成器 — Markdown + 图片 → docx
用法:
    python3 gen-report.py report.md -o output.docx
    python3 gen-report.py report.md -o output.docx --images docs/report/images/

Markdown 格式约定:
    # 标题          → 论文大标题（居中，二号宋体加粗）
    作者：xxx        → 作者行（居中）
    摘 要：xxx       → 摘要段落
    关键词：xxx      → 关键词段落
    ## 一、引言      → 一级标题（四号黑体加粗）
    ### （一）xxx    → 二级标题（四号宋体加粗）
    #### 1．xxx     → 三级标题（四号宋体）
    ##### xxx       → 四级标题（四号宋体加粗）
    正文段落         → 四号宋体，首行缩进1cm，1.5倍行距
    ![图注](path)   → 插入图片 + 图注（小四宋体居中）
    | 表头 | ...    → 表格
    [1] xxx         → 参考文献
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ─── 格式常量 ─────────────────────────────────────────

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEADING2_CN = '黑体'

PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3.2)
MARGIN_RIGHT = Cm(3.2)

TITLE_SIZE = Pt(22)       # 二号
NORMAL_SIZE = Pt(14)      # 四号 (≈14pt)
CAPTION_SIZE = Pt(12)     # 小四
SMALL_SIZE = Pt(10.5)     # 五号


# ─── 辅助函数 ─────────────────────────────────────────

def set_run_font(run, cn_font=FONT_CN, en_font=FONT_EN, size=NORMAL_SIZE, bold=False):
    """设置 run 的中英文字体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    # 设置东亚字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), en_font)
    rfonts.set(qn('w:hAnsi'), en_font)
    rfonts.set(qn('w:eastAsia'), cn_font)


def set_paragraph_format(paragraph, align=None, line_spacing=1.5, first_indent=None,
                         space_before=None, space_after=None):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    if align is not None:
        pf.alignment = align
    if line_spacing:
        pf.line_spacing = line_spacing
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after


def add_styled_paragraph(doc, text, cn_font=FONT_CN, en_font=FONT_EN, size=NORMAL_SIZE,
                         bold=False, align=None, line_spacing=1.5, first_indent=None,
                         space_before=None, space_after=None):
    """添加带完整样式的段落"""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=align, line_spacing=line_spacing,
                         first_indent=first_indent, space_before=space_before,
                         space_after=space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, cn_font=cn_font, en_font=en_font, size=size, bold=bold)
    return p


# ─── Markdown 解析 ─────────────────────────────────────

def parse_markdown(md_text, images_dir=None):
    """解析 Markdown 文本，返回结构化元素列表"""
    elements = []
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,5})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            elements.append(('heading', level, heading_match.group(2).strip()))
            i += 1
            continue

        # 图片
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            if images_dir:
                full_path = Path(images_dir) / Path(img_path).name
                if not full_path.exists():
                    full_path = Path(img_path)
            else:
                full_path = Path(img_path)
            elements.append(('image', str(full_path), caption))
            i += 1
            continue

        # 表格
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            elements.append(('table', table_lines))
            continue

        # 参考文献
        ref_match = re.match(r'^\[(\d+)\]\s*(.+)$', stripped)
        if ref_match:
            elements.append(('reference', ref_match.group(1), ref_match.group(2)))
            i += 1
            continue

        # 普通段落（合并连续非空行）
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if not next_stripped:
                break
            if re.match(r'^#{1,5}\s', next_stripped):
                break
            if re.match(r'^!\[', next_stripped):
                break
            if next_stripped.startswith('|'):
                break
            if re.match(r'^\[\d+\]', next_stripped):
                break
            para_lines.append(next_stripped)
            i += 1

        text = ' '.join(para_lines)
        # 检测特殊段落类型
        if text.startswith('作者') or text.startswith('Author'):
            elements.append(('author', text))
        elif text.startswith('摘') and '要' in text[:4]:
            elements.append(('abstract', text))
        elif text.startswith('关键词') or text.startswith('Keywords'):
            elements.append(('keywords', text))
        elif text.upper().startswith('ABSTRACT'):
            elements.append(('abstract_en', text))
        elif text.startswith('Key words') or text.upper().startswith('KEY WORDS'):
            elements.append(('keywords_en', text))
        else:
            elements.append(('paragraph', text))

    return elements


# ─── 表格解析 ─────────────────────────────────────────

def parse_table_lines(table_lines):
    """解析 Markdown 表格行"""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        # 跳过分隔行
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        rows.append(cells)
    return rows


# ─── 文档生成 ─────────────────────────────────────────

def generate_docx(elements, output_path, images_dir=None):
    """根据解析后的元素生成 docx"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    for elem in elements:
        etype = elem[0]

        if etype == 'heading':
            level = elem[1]
            text = elem[2]

            if level == 1:
                # 大标题：二号宋体加粗居中
                add_styled_paragraph(doc, text, size=TITLE_SIZE, bold=True,
                                     align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0,
                                     space_before=Pt(12), space_after=Pt(12))
            elif level == 2:
                # 一级标题：四号黑体加粗
                add_styled_paragraph(doc, text, cn_font=FONT_HEADING2_CN,
                                     size=NORMAL_SIZE, bold=True, line_spacing=1.5,
                                     space_before=Pt(12), space_after=Pt(12))
            elif level == 3:
                # 二级标题：四号宋体加粗
                add_styled_paragraph(doc, text, size=NORMAL_SIZE, bold=True,
                                     line_spacing=1.0, space_before=Pt(12),
                                     space_after=Pt(12))
            elif level == 4:
                # 三级标题：四号宋体
                add_styled_paragraph(doc, text, size=NORMAL_SIZE, bold=False,
                                     line_spacing=1.5, first_indent=Cm(1.0),
                                     space_before=Pt(6), space_after=Pt(6))
            elif level == 5:
                # 四级标题：四号宋体加粗
                add_styled_paragraph(doc, text, size=NORMAL_SIZE, bold=True,
                                     line_spacing=1.55, first_indent=Cm(1.0),
                                     space_before=Pt(6), space_after=Pt(6))

        elif etype == 'author':
            add_styled_paragraph(doc, elem[1], size=NORMAL_SIZE, bold=True,
                                 align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

        elif etype in ('abstract', 'abstract_en'):
            add_styled_paragraph(doc, elem[1], size=NORMAL_SIZE, bold=True,
                                 line_spacing=1.5, first_indent=Cm(1.0))

        elif etype in ('keywords', 'keywords_en'):
            add_styled_paragraph(doc, elem[1], size=NORMAL_SIZE, bold=True,
                                 line_spacing=1.5, first_indent=Cm(1.0))

        elif etype == 'paragraph':
            add_styled_paragraph(doc, elem[1], size=NORMAL_SIZE,
                                 align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                 line_spacing=1.5, first_indent=Cm(1.0))

        elif etype == 'image':
            img_path, caption = elem[1], elem[2]
            # 插入图片
            if Path(img_path).exists():
                p = doc.add_paragraph()
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
                run = p.add_run()
                # 图片宽度限制在版心宽度内
                max_width = PAGE_WIDTH - MARGIN_LEFT - MARGIN_RIGHT
                run.add_picture(img_path, width=min(Cm(14), max_width))
            else:
                add_styled_paragraph(doc, f'[图片缺失: {img_path}]',
                                     size=CAPTION_SIZE,
                                     align=WD_ALIGN_PARAGRAPH.CENTER,
                                     line_spacing=1.0)
            # 图注
            if caption:
                add_styled_paragraph(doc, caption, size=CAPTION_SIZE,
                                     align=WD_ALIGN_PARAGRAPH.CENTER,
                                     line_spacing=1.0, space_after=Pt(6))

        elif etype == 'table':
            rows = parse_table_lines(elem[1])
            if not rows:
                continue
            n_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < n_cols:
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                                             line_spacing=1.0)
                        run = p.add_run(cell_text)
                        is_header = (ri == 0)
                        set_run_font(run, size=SMALL_SIZE, bold=is_header)

            # 表后空行
            doc.add_paragraph()

        elif etype == 'reference':
            num, text = elem[1], elem[2]
            add_styled_paragraph(doc, f'[{num}] {text}', size=SMALL_SIZE,
                                 line_spacing=1.5, first_indent=Cm(0))

    doc.save(output_path)
    return output_path


# ─── 主入口 ─────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='研究报告生成器：Markdown → docx')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', default='report.docx', help='输出 docx 路径（默认 report.docx）')
    parser.add_argument('--images', default=None, help='图片目录（可选，用于解析图片相对路径）')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'错误：文件不存在 {input_path}')
        sys.exit(1)

    md_text = input_path.read_text(encoding='utf-8')
    elements = parse_markdown(md_text, images_dir=args.images)

    output = generate_docx(elements, args.output, images_dir=args.images)
    print(f'报告已生成: {output}')
    print(f'  段落数: {len([e for e in elements if e[0] == "paragraph"])}')
    print(f'  图片数: {len([e for e in elements if e[0] == "image"])}')
    print(f'  表格数: {len([e for e in elements if e[0] == "table"])}')
    print(f'  参考文献: {len([e for e in elements if e[0] == "reference"])}')


if __name__ == '__main__':
    main()
